"""
周报 Word 文档生成器 - python-docx
位置：D:\\demo_plan\\backend\\app\\services\\report_word_generator.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any
import base64
from io import BytesIO
from datetime import datetime

# 当前python-docx-valutico不支持标准图表导入，使用图片方式
CHART_SUPPORTED = False
print("ℹ️ 使用matplotlib生成高质量图表图片")


def create_word_report(data: Dict[str, Any]) -> bytes:
    """
    生成 Word 格式的周报

    Args:
        data: 报告数据字典（与 Markdown 模板相同）

    Returns:
        Word 文档的二进制数据
    """
    doc = Document()

    # 设置全局样式
    _set_global_styles(doc)

    # 1. 标题
    _add_title(doc, data['title'])

    # 2. 警告信息
    _add_warning_box(doc, data.get('warning_message', ''))

    # 3. 基本信息
    _add_basic_info(doc, data)

    # 4. 执行摘要
    _add_executive_summary(doc, data)

    # 5. 工作概览
    _add_work_overview(doc, data)

    # 6. 每日完成趋势
    _add_daily_trend(doc, data)

    # 7. 任务分析
    _add_task_analysis(doc, data)

    # 8. 习惯追踪
    _add_habit_tracking(doc, data)

    # 9. 个性化建议
    _add_suggestions(doc, data)

    # 10. 下周计划
    _add_next_week_plan(doc, data)

    # 保存为二进制
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.read()


def _set_global_styles(doc: Document):
    """设置全局样式"""
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)

    # 中文字体支持
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def _add_title(doc: Document, title: str):
    """添加标题"""
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置标题样式
    run = heading.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(33, 37, 41)

    doc.add_paragraph()  # 空行


def _add_warning_box(doc: Document, warning: str):
    """添加警告框"""
    if not warning:
        return

    # 创建带背景色的段落（使用表格模拟边框效果）
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    cell = table.rows[0].cells[0]
    cell.text = warning
    
    # 设置单元格样式
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(220, 53, 69)
    run.font.bold = True
    
    # 设置浅红色背景（通过 shading）
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'FFF2F0')  # 浅红色背景
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _add_basic_info(doc: Document, data: Dict[str, Any]):
    """添加基本信息"""
    info_text = f"报告周期：{data['start_date']} 至 {data['end_date']}\n"
    info_text += f"生成时间：{data['generate_time']}\n"
    info_text += f"用户：{data['user_nickname']} | 助手：{data['assistant_nickname']}"

    paragraph = doc.add_paragraph(info_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)

    # 添加分隔线（使用水平线）
    separator = doc.add_paragraph()
    separator.paragraph_format.space_before = Pt(6)
    separator.paragraph_format.space_after = Pt(6)
    run = separator.add_run('_' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(200, 200, 200)


def _add_executive_summary(doc: Document, data: Dict[str, Any]):
    """添加执行摘要"""
    doc.add_heading('📊 执行摘要', level=1)

    # 本周亮点
    doc.add_heading('✨ 本周亮点', level=2)
    for highlight in data.get('highlights', []):
        p = doc.add_paragraph(highlight, style='List Bullet')
        p.runs[0].font.size = Pt(10.5)

    # 关键问题
    doc.add_heading('⚠️ 关键问题', level=2)
    for issue in data.get('issues', []):
        p = doc.add_paragraph(issue, style='List Bullet')
        p.runs[0].font.size = Pt(10.5)

    # 改进方向
    doc.add_heading('💡 改进方向', level=2)
    for improvement in data.get('improvements', []):
        p = doc.add_paragraph(improvement, style='List Bullet')
        p.runs[0].font.size = Pt(10.5)

    doc.add_paragraph()


def _add_work_overview(doc: Document, data: Dict[str, Any]):
    """添加工作概览"""
    doc.add_heading('一、工作概览', level=1)

    # 核心数据表格 - 使用柔和样式
    doc.add_heading('1.1 核心数据', level=2)

    stats = data['stats']
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light List Accent 1'  # 使用柔和的表格样式
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格总宽度为更小（约12cm）
    for row in table.rows:
        for cell in row.cells:
            # 设置单元格宽度
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)

    # 表头
    headers = ['指标', '数值', '说明']
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 数据行
    rows_data = [
        ('总任务数', f"{stats['total']} 个", "本周所有任务"),
        ('已完成', f"{stats['completed']} 个 ({stats['completed_rate'] * 100:.1f}%)", "成功完成的任务"),
        ('进行中', f"{stats['in_progress']} 个 ({stats['in_progress'] / stats['total'] * 100:.1f}%)" if stats[
                                                                                                            'total'] > 0 else "0 个 (0%)",
         "待完成的任务"),
        ('已超时', f"{stats['overdue']} 个 ({stats['overtime_rate'] * 100:.1f}%)", "超过截止时间的任务"),
        ('总工作时长', f"{stats['total_hours']} 小时", "预计总工作时间"),
        ('日均任务数', f"{stats['daily_avg']} 个", "平均每天任务数"),
    ]

    for row_idx, (metric, value, description) in enumerate(rows_data, 1):
        table.rows[row_idx].cells[0].text = metric
        table.rows[row_idx].cells[1].text = value
        table.rows[row_idx].cells[2].text = description

        # 居中对齐
        for col_idx in range(3):
            table.rows[row_idx].cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 任务状态分布图 - 使用高质量matplotlib图片（尺寸优化）
    doc.add_heading('1.2 任务状态分布', level=2)
    _add_chart_image(doc, data['charts'].get('status_pie', ''), '任务状态分布图')
    
    # 添加图表说明文字
    stats = data['stats']
    summary_text = f"本周共处理 {stats['total']} 个任务，其中已完成 {stats['completed']} 个（{stats['completed_rate']*100:.1f}%），"
    summary_text += f"进行中 {stats['in_progress']} 个，已超时 {stats['overdue']} 个（{stats['overtime_rate']*100:.1f}%）。"
    if stats['overtime_rate'] > 0.4:
        summary_text += " 超时率偏高，建议下周优化任务分配。"
    else:
        summary_text += " 整体完成情况良好。"
    
    summary_para = doc.add_paragraph(summary_text)
    summary_para.paragraph_format.space_before = Pt(6)
    summary_para.paragraph_format.space_after = Pt(12)
    summary_para.runs[0].font.size = Pt(10)
    summary_para.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph()


def _add_daily_trend(doc: Document, data: Dict[str, Any]):
    """添加每日完成趋势"""
    doc.add_heading('二、每日完成趋势', level=1)

    # 趋势图表
    doc.add_heading('2.1 趋势图表', level=2)
    _add_chart_image(doc, data['charts'].get('daily_bar', ''), '每日完成任务趋势图')
    doc.add_paragraph()

    # 详细数据表格 - 使用柔和样式
    doc.add_heading('2.2 详细数据', level=2)

    daily_data = data.get('daily_data', [])
    if daily_data:
        table = doc.add_table(rows=len(daily_data) + 1, cols=5)
        table.style = 'Light List Accent 1'  # 使用柔和的表格样式
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格宽度
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)

        # 表头
        headers = ['日期', '星期', '完成任务', '总任务', '完成率']
        for col_idx, header in enumerate(headers):
            cell = table.rows[0].cells[col_idx]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 数据行
        for row_idx, day in enumerate(daily_data, 1):
            completion_rate = (day['completed'] / day['total'] * 100) if day['total'] > 0 else 0

            table.rows[row_idx].cells[0].text = day['date']
            table.rows[row_idx].cells[1].text = day['day']
            table.rows[row_idx].cells[2].text = str(day['completed'])
            table.rows[row_idx].cells[3].text = str(day['total'])
            table.rows[row_idx].cells[4].text = f"{completion_rate:.1f}%"

            # 居中对齐
            for col_idx in range(5):
                table.rows[row_idx].cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def _add_task_analysis(doc: Document, data: Dict[str, Any]):
    """添加任务分析"""
    doc.add_heading('三、任务分析', level=1)

    # 优先级分布
    doc.add_heading('3.1 优先级分布', level=2)
    priority_map = {'high': '高优先级', 'medium': '中优先级', 'low': '低优先级'}
    for priority, count in data['stats'].get('priority_distribution', {}).items():
        p = doc.add_paragraph(f"{priority_map.get(priority, priority)}: {count} 个", style='List Bullet')
        p.runs[0].font.size = Pt(10.5)

    # 任务类型分布
    doc.add_heading('3.2 任务类型分布', level=2)
    for type_name, count in data['stats'].get('type_distribution', {}).items():
        p = doc.add_paragraph(f"{type_name}: {count} 个", style='List Bullet')
        p.runs[0].font.size = Pt(10.5)

    doc.add_paragraph()


def _add_habit_tracking(doc: Document, data: Dict[str, Any]):
    """添加习惯追踪"""
    doc.add_heading('四、习惯追踪', level=1)

    habits = data.get('habits', {})
    habit_completion = habits.get('habit_completion', {})

    doc.add_heading('4.1 习惯完成情况', level=2)

    if habit_completion:
        habit_map = {'exercise': '运动', 'reading': '阅读', 'early_rise': '早起', 'study': '学习'}
        for habit, habit_data in habit_completion.items():
            text = f"{habit_map.get(habit, habit)}: {habit_data['completed']}/{habit_data['total']} ({habit_data['rate'] * 100:.0f}%)"
            p = doc.add_paragraph(text, style='List Bullet')
            p.runs[0].font.size = Pt(10.5)
    else:
        doc.add_paragraph('本周暂无习惯任务记录')

    # 一致性评分
    doc.add_heading('4.2 一致性评分', level=2)
    consistency_score = habits.get('consistency_score', 0)
    doc.add_paragraph(f"评分: {consistency_score * 100:.0f}%")
    doc.add_paragraph("说明: 基于按时完成率计算")

    doc.add_paragraph()


def _add_suggestions(doc: Document, data: Dict[str, Any]):
    """添加个性化建议"""
    doc.add_heading('五、个性化建议', level=1)

    llm_suggestions = data.get('llm_suggestions', '')
    
    if not llm_suggestions:
        doc.add_paragraph('暂无个性化建议')
        return

    # 处理 LLM 返回的 Markdown 格式建议
    lines = llm_suggestions.split('\n')
    for line in lines:
        line = line.strip()
        if line:
            # 如果是编号列表
            if len(line) > 2 and line[0:1].isdigit() and line[1:3] in ['. ', '、']:
                p = doc.add_paragraph(line[3:], style='List Number')
                p.runs[0].font.size = Pt(10.5)
            else:
                p = doc.add_paragraph(line)
                p.runs[0].font.size = Pt(10.5)

    doc.add_paragraph()


def _add_next_week_plan(doc: Document, data: Dict[str, Any]):
    """添加下周计划"""
    doc.add_heading('六、下周计划建议', level=1)

    stats = data['stats']

    # 待完成任务
    doc.add_heading('6.1 待完成任务', level=2)
    if stats.get('overdue', 0) > 0:
        doc.add_paragraph(f"⚠️ 有 {stats['overdue']} 个任务已超时，建议优先处理：")
    else:
        doc.add_paragraph("✅ 无超时任务，继续保持！")

    # 负载建议
    doc.add_heading('6.2 负载建议', level=2)
    overtime_rate = stats.get('overtime_rate', 0)

    if overtime_rate > 0.6:
        doc.add_paragraph(f"负载警告: 本周超时率 {overtime_rate * 100:.0f}% 过高")
        doc.add_paragraph("建议: 下周任务量减少 30%，聚焦高优先级任务")
    elif overtime_rate > 0.4:
        doc.add_paragraph(f"负载提示: 本周超时率 {overtime_rate * 100:.0f}%")
        doc.add_paragraph("建议: 合理分配时间，避免任务堆积")
    else:
        doc.add_paragraph(f"负载良好: 本周超时率 {overtime_rate * 100:.0f}%")
        doc.add_paragraph("建议: 继续保持当前节奏")

    doc.add_paragraph()

    # 页脚
    footer_text = f"报告生成: {data['assistant_nickname']} · 智能助手\n下次报告: {data['next_report_date']}"
    footer = doc.add_paragraph(footer_text)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(20)

    run = footer.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True


def _add_native_chart(doc: Document, stats: Dict, alt_text: str):
    """添加原生Word图表（环形图/饼图）"""
    try:
        # 准备数据
        completed = stats.get('completed', 0)
        in_progress = stats.get('in_progress', 0)
        overdue = stats.get('overdue', 0)
        
        if completed + in_progress + overdue == 0:
            doc.add_paragraph(f"[图表: {alt_text} - 无数据]")
            return
        
        # 创建图表数据
        chart_data = CategoryChartData()
        chart_data.categories = []
        chart_data.add_series('任务状态')
        
        if completed > 0:
            chart_data.categories.append('已完成')
            chart_data.series[0].add_value(completed)
        if in_progress > 0:
            chart_data.categories.append('进行中')
            chart_data.series[0].add_value(in_progress)
        if overdue > 0:
            chart_data.categories.append('已超时')
            chart_data.series[0].add_value(overdue)
        
        # 添加图表到文档（宽度11.5cm，高度7cm）
        graphic_frame = doc.add_chart(
            XL_CHART_TYPE.PIE,
            Cm(11.5),
            Cm(7),
            chart_data
        )
        
        # 获取图表对象
        chart = graphic_frame.chart
        
        # 添加图表标题
        chart.has_title = True
        chart.chart_title.text_frame.text = alt_text
        
        # 添加图例
        chart.has_legend = True
        chart.legend.position = XL_LEGEND_POSITION.BOTTOM
        chart.legend.include_in_layout = False
        
        # 添加数据标签（百分比）
        plot = chart.plots[0]
        plot.has_data_labels = True
        data_labels = plot.data_labels
        data_labels.show_percent = True
        data_labels.show_category_name = False
        data_labels.number_format = '0.0%'
        data_labels.position = XL_LABEL_POSITION.INSIDE_END
        
        doc.add_paragraph()
        
    except Exception as e:
        print(f"⚠️ 添加原生图表失败: {e}")
        import traceback
        traceback.print_exc()
        # 降级为图片显示
        doc.add_paragraph(f"[图表加载失败: {alt_text}]")


def _add_chart_image(doc: Document, base64_data: str, alt_text: str):
    """添加图表图片（用于柱状图和环形图）"""
    if not base64_data or not base64_data.startswith('data:image'):
        doc.add_paragraph(f"[图片: {alt_text}]")
        return

    try:
        # 提取 Base64 数据
        header, encoded = base64_data.split(',', 1)
        image_data = base64.b64decode(encoded)

        # 添加到文档 - 使用更小的尺寸（宽度8cm，更紧凑）
        image_stream = BytesIO(image_data)
        doc.add_picture(image_stream, width=Cm(8))

        # 添加图片说明
        caption = doc.add_paragraph(alt_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(9)
        caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    except Exception as e:
        print(f"️ 添加图片失败: {e}")
        doc.add_paragraph(f"[图片加载失败: {alt_text}]")
